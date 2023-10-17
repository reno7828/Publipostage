import os
import shutil
from tkinter import *
from tkinter import filedialog
from docx import Document
import pandas as pd
import zipfile
from io import BytesIO

from docx import Document

def process(doc_path):
    # Ouvre le document Word
    doc = Document(doc_path)

    # Initialise le texte
    text = ""

    # Parcours chaque paragraphe du document et ajoute son texte
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'

    return text


# Reste du code inchangé...

def generate_documents():
    excel_path = filedialog.askopenfilename(title="Sélectionner le fichier Excel")
    doc_path = filedialog.askopenfilename(title="Sélectionner le modèle Word")
    colonnes_str = entry_colonnes.get()
    colonnes = [colonne.strip() for colonne in colonnes_str.split(',')]

    dataframe = pd.read_excel(excel_path)
    noms_distincts = dataframe[colonnes[0]].unique()

    temp_dir = os.path.join(os.getcwd(), 'temp_dir')
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    for nom in noms_distincts:
        text = process(doc_path)  # Convertir le document Word en texte brut
        for colonne in colonnes:
            if colonne.lower() == 'pour':
                # Formater le pourcentage sans séparateur de milliers et ajouter le symbole %
                valeur = str(float(dataframe.loc[dataframe[colonnes[0]] == nom, colonne].iloc[0]) * 100)
                # Supprimer les zéros inutiles après la virgule
                valeur = valeur.rstrip('0').rstrip('.') + '%'
            else:
                valeur = str(dataframe.loc[dataframe[colonnes[0]] == nom, colonne].iloc[0])
            text = text.replace(colonne, valeur)

        document_personnalise = Document()
        document_personnalise.add_paragraph(text)  # Ajouter le texte au document personnalisé
        docx_file_path = os.path.join(temp_dir, f"test_{nom}.docx")
        document_personnalise.save(docx_file_path)

    zip_buffer = BytesIO()  # Créer un buffer pour le fichier ZIP
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for docx_file in os.listdir(temp_dir):
            docx_file_path = os.path.join(temp_dir, docx_file)
            zip_file.write(docx_file_path, os.path.basename(docx_file_path))  # Ajouter le fichier DOCX au fichier ZIP

    # Supprimer le répertoire temporaire
    shutil.rmtree(temp_dir)

    # Définir l'emplacement où enregistrer le fichier ZIP
    zip_path = filedialog.asksaveasfilename(defaultextension=".zip", filetypes=[("Fichiers ZIP", "*.zip")])

    if zip_path:
        with open(zip_path, 'wb') as zip_file:
            zip_file.write(zip_buffer.getvalue())

# Créer la fenêtre tkinter
root = Tk()
root.title("Génération de lettres personnalisées")

# Interface utilisateur
label_colonnes = Label(root, text="Colonnes (séparées par des virgules):")
label_colonnes.pack(pady=10)
entry_colonnes = Entry(root, width=50)
entry_colonnes.pack(pady=5)
entry_colonnes.insert(0, "NOM,MAMO")  # Exemple de valeurs initiales

# Définir la taille de la fenêtre (largeur x hauteur)
root.geometry("600x400")

generate_button = Button(root, text="Générer les lettres", command=generate_documents)
generate_button.pack(pady=20)

# Lancer la boucle principale de tkinter
root.mainloop()
